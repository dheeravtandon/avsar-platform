"""Word-document helpers for the AVSAR document set.

Gives every .docx the same cover block, heading scale, table style and footer,
so the four narrative documents read as one set rather than four drafts.
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from common import CLASSIFICATION, DOC_DATE, DOC_IDS, PREPARED_BY, PROJECT_LONG, VERSION

NAVY = RGBColor(0x0B, 0x24, 0x47)
ACCENT = RGBColor(0x05, 0x63, 0x9E)
INK = RGBColor(0x34, 0x40, 0x54)
MUTED = RGBColor(0x66, 0x70, 0x85)
SAFFRON = RGBColor(0xC2, 0x6A, 0x0D)


def new_document():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, colour, before, after in [
        ("Heading 1", 17, NAVY, 20, 8),
        ("Heading 2", 13, NAVY, 16, 6),
        ("Heading 3", 11.5, ACCENT, 12, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    return doc


def shade(cell, hex_colour):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(el)


def cell_text(cell, text, bold=False, size=9.5, colour=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align:
        p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = colour
    return cell


def add_table(doc, headers, rows, widths=None, header_fill="0B2447", font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, head in enumerate(headers):
        c = table.rows[0].cells[i]
        cell_text(c, head, bold=True, size=font_size, colour=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, header_fill)

    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=font_size)
            if r % 2 == 1:
                shade(cells[i], "F9FAFB")

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def cover(doc, key, subtitle, purpose, revisions, approvals=None):
    """Standard cover page: title block, document control table, purpose, revisions."""
    doc_id, title = DOC_IDS[key]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(PROJECT_LONG)
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    run = p.add_run("Document control")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY

    related = ", ".join(v[0] for k, v in DOC_IDS.items() if k != key)
    add_table(doc, ["Field", "Value"], [
        ["Document ID", doc_id],
        ["Version", VERSION],
        ["Date", DOC_DATE],
        ["Prepared by", PREPARED_BY],
        ["Status", "Baselined"],
        ["Related documents", related],
        ["Classification", CLASSIFICATION],
    ], widths=[1.8, 4.8])

    doc.add_heading("Purpose of this document", level=2)
    doc.add_paragraph(purpose)

    doc.add_heading("Revision history", level=2)
    add_table(doc, ["Version", "Date", "Author", "Change"], revisions,
              widths=[0.8, 1.1, 1.1, 3.6])

    if approvals:
        doc.add_heading("Approval", level=2)
        add_table(doc, ["Role", "Name", "Signature", "Date"], approvals,
                  widths=[2.0, 1.8, 1.5, 1.1])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(f"{lead} ")
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def callout(doc, label, text, colour="F2F4F7"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = INK
    shade(cell, colour)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def footer_note(doc, text):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.font.name = "Calibri"
